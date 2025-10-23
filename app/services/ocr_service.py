"""
OCR servisi - PDF ve görsel belgelerden metin çıkarma
"""
import logging
from pathlib import Path
from typing import Optional
import warnings
import pypdf

# PyPDF/PyMuPDF uyarılarını bastır (hatalı PDF formatları için)
warnings.filterwarnings("ignore", message="Multiple definitions in dictionary")
warnings.filterwarnings("ignore", category=pypdf.errors.PdfReadWarning)

logger = logging.getLogger(__name__)


class OCRService:
    """OCR ve metin çıkarma servisi"""

    def __init__(self):
        self.supported_formats = {'.pdf', '.docx', '.doc', '.jpg', '.jpeg', '.png'}

    def extract_text_from_pdf(self, file_path: Path) -> str:
        """
        PDF'den metin çıkar

        Args:
            file_path: PDF dosya yolu

        Returns:
            Çıkarılan metin
        """
        try:
            text_content = []

            with open(file_path, 'rb') as f:
                pdf_reader = pypdf.PdfReader(f)
                num_pages = len(pdf_reader.pages)

                logger.info(f"PDF okunuyor: {num_pages} sayfa")

                for page_num in range(num_pages):
                    page = pdf_reader.pages[page_num]
                    text = page.extract_text()

                    if text.strip():
                        text_content.append(f"\n=== Sayfa {page_num + 1} ===\n")
                        text_content.append(text)

            full_text = '\n'.join(text_content)
            logger.info(f"PDF'den {len(full_text)} karakter metin çıkarıldı")

            return full_text

        except Exception as e:
            logger.error(f"PDF okuma hatası: {str(e)}")
            raise

    def extract_text_from_docx(self, file_path: Path) -> str:
        """
        DOCX'ten metin çıkar

        Args:
            file_path: DOCX dosya yolu

        Returns:
            Çıkarılan metin
        """
        try:
            from docx import Document

            doc = Document(file_path)
            text_content = []

            # Paragrafları al
            for para in doc.paragraphs:
                if para.text.strip():
                    text_content.append(para.text)

            # Tabloları al
            for table in doc.tables:
                for row in table.rows:
                    row_text = ' | '.join(cell.text for cell in row.cells)
                    if row_text.strip():
                        text_content.append(row_text)

            full_text = '\n'.join(text_content)
            logger.info(f"DOCX'ten {len(full_text)} karakter metin çıkarıldı")

            return full_text

        except Exception as e:
            logger.error(f"DOCX okuma hatası: {str(e)}")
            # python-docx yoksa fallback
            logger.warning("python-docx bulunamadı, boş metin dönüyorum")
            return ""

    def extract_text(self, file_path: Path) -> str:
        """
        Dosya tipine göre metin çıkar

        Args:
            file_path: Dosya yolu

        Returns:
            Çıkarılan metin
        """
        extension = file_path.suffix.lower()

        if extension == '.pdf':
            return self.extract_text_from_pdf(file_path)
        elif extension in ['.docx', '.doc']:
            return self.extract_text_from_docx(file_path)
        elif extension in ['.jpg', '.jpeg', '.png', '.bmp']:
            # OCR gerekli (EasyOCR)
            return self.extract_text_from_image(file_path)
        else:
            logger.warning(f"Desteklenmeyen format: {extension}")
            return ""

    def extract_text_from_image(self, file_path: Path) -> str:
        """
        Görsel dosyadan OCR ile metin çıkar

        Not: EasyOCR kurulumu gerekiyor
        """
        try:
            import easyocr

            reader = easyocr.Reader(['tr', 'en'], gpu=False)
            logger.info(f"OCR başlatılıyor: {file_path}")

            result = reader.readtext(str(file_path))

            text_content = []
            for detection in result:
                text = detection[1]
                text_content.append(text)

            full_text = '\n'.join(text_content)
            logger.info(f"OCR'den {len(full_text)} karakter metin çıkarıldı")

            return full_text

        except ImportError:
            logger.warning("EasyOCR kurulu değil, görsel OCR yapılamıyor")
            return ""
        except Exception as e:
            logger.error(f"OCR hatası: {str(e)}")
            return ""

    def clean_text(self, text: str) -> str:
        """
        Metni temizle (fazla boşluklar, satır sonları vs.)

        Args:
            text: Ham metin

        Returns:
            Temizlenmiş metin
        """
        # Çoklu boşlukları tek boşluğa çevir
        import re
        text = re.sub(r'\s+', ' ', text)

        # Satır başı/sonu boşlukları temizle
        text = text.strip()

        return text
